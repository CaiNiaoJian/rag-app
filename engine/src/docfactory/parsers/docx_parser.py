"""docx 解析器（python-docx，MIT；03 章 §1 直解路径）。

覆盖：标题层级 → section 树、段落、列表（numPr）、表格（gridSpan/vMerge → colspan/rowspan）、
内嵌图片落 assets、页眉页脚独立节点。

**关于页码的取舍（重要）**：docx 是流式文档，页码由渲染器（Word/LibreOffice）在排版时才产生，
文件里根本不存在「第几页」这个事实。可用的只有两类线索：
  ① 显式分页符 ``w:br[@w:type="page"]`` 与 ``w:pageBreakBefore``（作者主动插入的）；
  ② ``w:lastRenderedPageBreak`` —— Word 保存时写入的「上次渲染的分页位置」，相当准，
     但 LibreOffice 转换出来的 docx 通常没有。
本解析器用 ①+② 累加估算 ``prov.page``：能拿到线索就分页，拿不到就整篇 page=1。
**宁可少估也不误导** —— 溯源里一个错误的页码比一个笼统的 page=1 更有害（用户会照着去翻页）。

降级：结构化解析抛异常时由上层调用 ``parse_text_fallback()``，退成整篇纯文本段落（L2）。
"""

from __future__ import annotations

import re
from collections.abc import Iterator
from pathlib import Path
from typing import Any

from docx import Document as open_docx
from docx.oxml.ns import qn

from docfactory.errors import DocFactoryError
from docfactory.ir import IRNode, NodeContent, Prov, TableCell, TableContent
from docfactory.parsers import ParseEnv, assert_ooxml_readable
from docfactory.taskspec import EVENT_PROGRESS

# 标题样式识别：英文/中文本地化样式名 + 样式 id 兜底（同一份文档可能三种都出现）
_HEADING_NAME = re.compile(r"^(?:heading|标题)\s*([1-9])$", re.IGNORECASE)
_HEADING_ID = re.compile(r"^heading([1-9])$", re.IGNORECASE)
_CAPTION_STYLE = re.compile(r"caption|题注|图注|表注", re.IGNORECASE)

# 进度上报粒度：docx 没有页的概念，用 body 块序号做进度分母（附 unit 标明口径）
_PROGRESS_EVERY = 25

_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff", ".emf", ".wmf", ".webp"}

# VML 图片元素的完整限定名。python-docx 的 qn() 只认识它自己 nsmap 里的前缀，
# 'v'（urn:schemas-microsoft-com:vml）不在其中，写 qn("v:imagedata") 会直接 KeyError。
_VML_IMAGEDATA = "{urn:schemas-microsoft-com:vml}imagedata"


def _open(src: Path) -> Any:
    assert_ooxml_readable(src)
    try:
        return open_docx(str(src))
    except DocFactoryError:
        raise
    except Exception as exc:
        raise DocFactoryError("E01", f"打开 docx 失败：{src.name}（{type(exc).__name__}: {exc}）") from exc


# ---------------------------------------------------------------- XML 小工具


def _para_text(p: Any) -> str:
    """段落纯文本：w:t 拼接 + w:tab 转空格 + w:br 转换行（不含域代码与批注）。"""
    parts: list[str] = []
    for el in p.iter():
        tag = el.tag
        if tag == qn("w:t"):
            parts.append(el.text or "")
        elif tag == qn("w:tab"):
            parts.append("\t")
        elif tag == qn("w:br") and el.get(qn("w:type")) != "page":
            parts.append("\n")
    return "".join(parts).strip()


def _style_name(p: Any, styles: dict[str, str]) -> str:
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        return ""
    style_el = pPr.find(qn("w:pStyle"))
    if style_el is None:
        return ""
    sid = style_el.get(qn("w:val")) or ""
    return styles.get(sid, sid)


def _heading_level(p: Any, styles: dict[str, str]) -> int | None:
    """返回 1..6 或 None。样式名 → 样式 id → outlineLvl 三重兜底。"""
    pPr = p.find(qn("w:pPr"))
    style_el = pPr.find(qn("w:pStyle")) if pPr is not None else None
    sid = (style_el.get(qn("w:val")) or "") if style_el is not None else ""
    name = styles.get(sid, sid)

    for pattern, text in ((_HEADING_NAME, name), (_HEADING_ID, sid)):
        m = pattern.match(text.strip())
        if m:
            return min(6, int(m.group(1)))
    if name.strip().lower() in ("title", "标题") or sid.lower() == "title":
        return 1
    if pPr is not None:
        lvl = pPr.find(qn("w:outlineLvl"))
        if lvl is not None:
            try:
                val = int(lvl.get(qn("w:val")) or "9")
            except ValueError:
                return None
            if 0 <= val <= 5:
                return val + 1
    return None


def _numpr_of(pPr: Any) -> tuple[str, int] | None:
    """从一个 pPr 元素里读 (numId, ilvl)。"""
    if pPr is None:
        return None
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None
    num_id_el = numPr.find(qn("w:numId"))
    if num_id_el is None:
        return None
    ilvl_el = numPr.find(qn("w:ilvl"))
    try:
        ilvl = int(ilvl_el.get(qn("w:val")) or "0") if ilvl_el is not None else 0
    except ValueError:
        ilvl = 0
    return (num_id_el.get(qn("w:val")) or "0", ilvl)


def _num_info(p: Any, style_nums: dict[str, tuple[str, int]]) -> tuple[str, int] | None:
    """列表信息 (numId, ilvl)；非列表段落返回 None。

    编号可以挂在段落上（``w:pPr/w:numPr``，Word 里手动加编号时如此），
    也可以挂在**样式**上（"List Number"/"List Bullet" 这类内置样式，段落本身没有 numPr）。
    只看段落会漏掉后者 —— 而后者恰恰是规范文档最常见的写法。
    """
    pPr = p.find(qn("w:pPr"))
    direct = _numpr_of(pPr)
    if direct is not None:
        return direct
    style_el = pPr.find(qn("w:pStyle")) if pPr is not None else None
    sid = (style_el.get(qn("w:val")) or "") if style_el is not None else ""
    return style_nums.get(sid)


def _page_advance(p: Any) -> int:
    """本段落引入了几次分页（估算依据见模块 docstring）。"""
    advance = 0
    pPr = p.find(qn("w:pPr"))
    if pPr is not None and pPr.find(qn("w:pageBreakBefore")) is not None:
        advance += 1
    for el in p.iter():
        rendered = el.tag == qn("w:lastRenderedPageBreak")
        explicit = el.tag == qn("w:br") and el.get(qn("w:type")) == "page"
        if rendered or explicit:
            advance += 1
    return advance


def _style_map(document: Any) -> dict[str, str]:
    """样式 id → 样式名（本地化名也在这里）。取不到就退回空表，靠 id 正则兜底。"""
    out: dict[str, str] = {}
    try:
        for style in document.styles.element.findall(qn("w:style")):
            sid = style.get(qn("w:styleId")) or ""
            name_el = style.find(qn("w:name"))
            if sid and name_el is not None:
                out[sid] = name_el.get(qn("w:val")) or ""
    except Exception:  # 样式表损坏不该拖垮整篇解析
        return {}
    return out


def _style_num_map(document: Any) -> dict[str, tuple[str, int]]:
    """样式 id → (numId, ilvl)：把「编号定义在样式里」的列表也认出来（见 _num_info）。"""
    out: dict[str, tuple[str, int]] = {}
    try:
        styles = document.styles.element.findall(qn("w:style"))
    except Exception:
        return out
    for style in styles:
        sid = style.get(qn("w:styleId")) or ""
        info = _numpr_of(style.find(qn("w:pPr")))
        if sid and info is not None:
            out[sid] = info
    return out


def _ordered_map(document: Any) -> dict[str, bool]:
    """numId → 是否有序列表（numFmt != bullet）。取不到时默认按有序处理。"""
    result: dict[str, bool] = {}
    try:
        numbering = document.part.numbering_part.element
    except Exception:
        return result
    abstract: dict[str, bool] = {}
    for el in numbering.findall(qn("w:abstractNum")):
        aid = el.get(qn("w:abstractNumId")) or ""
        fmt = ""
        lvl = el.find(qn("w:lvl"))
        if lvl is not None:
            fmt_el = lvl.find(qn("w:numFmt"))
            fmt = (fmt_el.get(qn("w:val")) or "") if fmt_el is not None else ""
        abstract[aid] = fmt.lower() != "bullet"
    for el in numbering.findall(qn("w:num")):
        nid = el.get(qn("w:numId")) or ""
        aid_el = el.find(qn("w:abstractNumId"))
        aid = (aid_el.get(qn("w:val")) or "") if aid_el is not None else ""
        result[nid] = abstract.get(aid, True)
    return result


def _iter_images(el: Any, document: Any) -> Iterator[tuple[bytes, str]]:
    """产出元素内所有内嵌图片的 (二进制, 扩展名)。

    同时覆盖 DrawingML（``a:blip``，现代 Word）与 VML（``v:imagedata``，老文档/转换产物）。
    """
    rels = document.part.related_parts
    seen: set[str] = set()
    for blip in el.iter(qn("a:blip")):
        rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        if rid and rid not in seen:
            seen.add(rid)
            yield from _image_from_rel(rels, rid)
    # VML 前缀 'v' 不在 python-docx 的 nsmap 里，qn() 会 KeyError；直接写完整 URI。
    # 这条路径踩中过一次真实事故：任何 docx 只要走到这里就整篇降级到 L2，标题与表格全丢。
    for image in el.iter(_VML_IMAGEDATA):
        rid = image.get(qn("r:id"))
        if rid and rid not in seen:
            seen.add(rid)
            yield from _image_from_rel(rels, rid)


def _image_from_rel(rels: Any, rid: str) -> Iterator[tuple[bytes, str]]:
    try:
        part = rels[rid]
        blob = part.blob
        ext = Path(str(part.partname)).suffix.lower() or ".png"
    except Exception:  # 关系缺失/外链图片：跳过，不影响正文
        return
    if ext in _IMAGE_EXTS and blob:
        yield blob, ext


# ---------------------------------------------------------------- 表格


def _tc_text(tc: Any) -> str:
    """单元格文本：只取直属段落，嵌套表格的文本不并入（避免层级串味）。"""
    lines = [_para_text(p) for p in tc.findall(qn("w:p"))]
    return "\n".join(line for line in lines if line).strip()


def _tc_span(tc: Any) -> tuple[int, str | None]:
    """返回 (colspan, vMerge 状态)；vMerge 状态为 "restart" / "continue" / None。"""
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        return 1, None
    colspan = 1
    grid = tcPr.find(qn("w:gridSpan"))
    if grid is not None:
        try:
            colspan = max(1, int(grid.get(qn("w:val")) or "1"))
        except ValueError:
            colspan = 1
    vmerge = tcPr.find(qn("w:vMerge"))
    if vmerge is None:
        return colspan, None
    val = (vmerge.get(qn("w:val")) or "continue").lower()
    return colspan, "restart" if val == "restart" else "continue"


def _table_content(tbl: Any) -> TableContent:
    """docx 表格 → IR cells。

    竖向合并的续格在 XML 里是**实际存在**的 ``w:tc``，所以逐行推进的列游标天然对齐；
    遇到 vMerge=continue 就把上方那个起始格的 rowspan +1，自己不生成新 cell
    （IR 约定合并区只保留左上格，见 ir.table_to_grid）。
    """
    cells: list[TableCell] = []
    open_merges: dict[int, TableCell] = {}   # 列号 → 正在纵向延伸的起始格
    rows = tbl.findall(qn("w:tr"))
    for r, tr in enumerate(rows):
        header_row = _is_header_row(tr) or r == 0
        col = 0
        for tc in tr.findall(qn("w:tc")):
            colspan, vmerge = _tc_span(tc)
            if vmerge == "continue":
                origin = open_merges.get(col)
                if origin is not None:
                    origin.rowspan += 1
                else:  # 起始格缺失（文档不规范）：当普通空格处理，保证网格完整
                    cells.append(TableCell(r=r, c=col, colspan=colspan, text=_tc_text(tc)))
                col += colspan
                continue
            cell = TableCell(
                r=r, c=col, colspan=colspan, text=_tc_text(tc), is_header=header_row
            )
            cells.append(cell)
            if vmerge == "restart":
                open_merges[col] = cell
            else:
                open_merges.pop(col, None)
            col += colspan
    return TableContent(cells=cells)


def _is_header_row(tr: Any) -> bool:
    trPr = tr.find(qn("w:trPr"))
    return trPr is not None and trPr.find(qn("w:tblHeader")) is not None


# ---------------------------------------------------------------- 主流程


class _Walker:
    """一次 docx 遍历的可变状态（章节栈、页码估算、列表聚合）。"""

    def __init__(self, document: Any, env: ParseEnv):
        self.document = document
        self.env = env
        self.styles = _style_map(document)
        self.style_nums = _style_num_map(document)
        self.ordered = _ordered_map(document)
        self.page = 1
        self.max_page = 1
        self.stack: list[tuple[int, IRNode]] = []      # (level, section 节点)
        self.list_node: IRNode | None = None
        self.list_key: tuple[str, int] | None = None
        self.last_figure: IRNode | None = None

    # ---- 节点归属 ----

    def parent(self) -> IRNode | None:
        return self.stack[-1][1] if self.stack else None

    def prov(self) -> list[Prov]:
        return [Prov(page=self.page)]

    def push_section(self, level: int, text: str) -> None:
        while self.stack and self.stack[-1][0] >= level:
            self.stack.pop()
        node = self.env.builder.add(
            "section", parent=self.parent(), level=level,
            content=NodeContent(text=text), prov=self.prov(),
        )
        self.stack.append((level, node))
        self.list_node = None

    def add_paragraph(self, text: str) -> IRNode:
        self.list_node = None
        return self.env.builder.add(
            "paragraph", parent=self.parent(), content=NodeContent(text=text), prov=self.prov()
        )

    def add_list_item(self, key: tuple[str, int], text: str) -> None:
        """连续同 numId 的段落聚合进同一个 list 节点（04 章 §2.1：list 的子节点是段落）。"""
        if self.list_node is None or self.list_key != key:
            self.list_node = self.env.builder.add(
                "list", parent=self.parent(),
                content=NodeContent(ordered=self.ordered.get(key[0], True)),
                prov=self.prov(),
            )
            self.list_key = key
        self.env.builder.add(
            "paragraph", parent=self.list_node,
            content=NodeContent(text=text), prov=self.prov(),
        )

    # ---- 块处理 ----

    def handle_paragraph(self, p: Any) -> None:
        self.page += _page_advance(p)
        self.max_page = max(self.max_page, self.page)

        for blob, ext in _iter_images(p, self.document):
            ref = self.env.save_asset(blob, ext)
            self.last_figure = self.env.builder.add(
                "figure", parent=self.parent(),
                content=NodeContent(image_ref=ref, ocr=False), prov=self.prov(),
            )
            self.list_node = None

        text = _para_text(p)
        if not text:
            return

        style = _style_name(p, self.styles)
        if self.last_figure is not None and _CAPTION_STYLE.search(style):
            # 题注样式紧跟图片：并入 figure.caption，不再单独成段
            self.last_figure.content.caption = text
            self.last_figure = None
            return

        level = _heading_level(p, self.styles)
        if level is not None:
            self.push_section(level, text)
            self.last_figure = None
            return

        num = _num_info(p, self.style_nums)
        if num is not None:
            self.add_list_item(num, text)
        else:
            self.add_paragraph(text)
        self.last_figure = None

    def handle_table(self, tbl: Any) -> None:
        self.page += sum(_page_advance(p) for p in tbl.iter(qn("w:p")))
        self.max_page = max(self.max_page, self.page)
        content = _table_content(tbl)
        if not content.cells:
            return
        self.env.builder.add(
            "table", parent=self.parent(),
            content=NodeContent(table=content), prov=self.prov(),
        )
        self.list_node = None
        self.last_figure = None


def parse(src: Path, env: ParseEnv) -> None:
    """结构化解析（L0 直解）：按 body 顺序遍历段落与表格，再补页眉页脚。"""
    document = _open(src)
    walker = _Walker(document, env)

    body = document.element.body
    blocks = [el for el in body.iterchildren() if el.tag in (qn("w:p"), qn("w:tbl"))]
    total = max(1, len(blocks))
    for idx, el in enumerate(blocks, start=1):
        if el.tag == qn("w:p"):
            walker.handle_paragraph(el)
        else:
            walker.handle_table(el)
        if idx % _PROGRESS_EVERY == 0 or idx == total:
            # unit=block：docx 无页概念，进度分母是正文块序号而非页码（见模块 docstring）
            env.check_cancel()
            if env.ctx is not None:
                env.ctx.progress(
                    EVENT_PROGRESS,
                    {"page": idx, "total": total, "stage": "parse", "unit": "block"},
                )

    _add_header_footer(document, env, walker.max_page)

    env.page_count = walker.max_page
    for page in range(1, walker.max_page + 1):
        env.mark_level(page, "L0")


def _add_header_footer(document: Any, env: ParseEnv, page: int) -> None:
    """页眉/页脚各自独立节点（04 章 §2.1，导出时可按 drop_header_footer 剔除）。

    同一段文字在多个 section 里重复出现是常态（整篇统一页眉），按文本去重只留一份。
    """
    seen: set[str] = set()
    for section in document.sections:
        for kind, part in (("header", section.header), ("footer", section.footer)):
            try:
                paragraphs = part.paragraphs
            except Exception:
                continue
            for p in paragraphs:
                text = _para_text(p._p)
                if not text or text in seen:
                    continue
                seen.add(text)
                env.builder.add(
                    kind, content=NodeContent(text=text), prov=[Prov(page=page)]  # type: ignore[arg-type]
                )


def parse_text_fallback(src: Path, env: ParseEnv) -> None:
    """L2 兜底：只抽段落与表格文本，放弃结构（03 章 §5.1 Office 两级降级）。"""
    document = _open(src)
    for el in document.element.body.iterchildren():
        if el.tag == qn("w:p"):
            text = _para_text(el)
            if text:
                env.builder.add("paragraph", content=NodeContent(text=text), prov=[Prov(page=1)])
        elif el.tag == qn("w:tbl"):
            lines = [
                "\t".join(_tc_text(tc) for tc in tr.findall(qn("w:tc")))
                for tr in el.findall(qn("w:tr"))
            ]
            text = "\n".join(line for line in lines if line.strip())
            if text:
                env.builder.add("paragraph", content=NodeContent(text=text), prov=[Prov(page=1)])
    env.page_count = 1
    env.mark_level(1, "L2")
