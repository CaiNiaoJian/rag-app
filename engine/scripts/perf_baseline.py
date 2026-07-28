"""性能基线实测（08 章 §4 的 M1 交付物「性能基线实测，修订 NFR 与体积表」）。

01 章 §2.2 的 NFR 数值是**拍出来的估计**，08 章 §3.2 明确写着「M1 实测后 M2 冻结为
验收门禁」。冻结之前得先有数——这个脚本就是那个数的来源。

它跑的是真实链路：拉起引擎子进程，走 HTTP 提交任务，按页轮询直到终态，
顺带采样进程内存。**不是**在进程内直接调函数——后者测不到 IPC、SSE、SQLite 写入
与 PyInstaller 冷启动这些真实开销，而用户体感恰恰由它们决定。

    cd engine && uv run python scripts/perf_baseline.py
    cd engine && uv run python scripts/perf_baseline.py --exe dist/engine/engine.exe
    cd engine && uv run python scripts/perf_baseline.py --report ../corpus/perf-baseline.md

退出码恒为 0：这是**测量工具不是门禁**。门禁值要等 M2 依据这里的数据冻结，
现在就拿实测值反过来当门禁，等于把「今天恰好这么快」写死成验收标准。
"""

from __future__ import annotations

import argparse
import json
import os
import secrets
import subprocess
import sys
import tempfile
import time
from contextlib import suppress
from dataclasses import dataclass, field
from pathlib import Path

import httpx

ROOT = Path(__file__).resolve().parents[2]
ENGINE_DIR = ROOT / "engine"

READY_TIMEOUT_S = 30.0
TASK_TIMEOUT_S = 900.0
MEM_SAMPLE_INTERVAL_S = 0.5

# 01 章 §2.2 / 08 章 §3.2 的目标值。这里只用于在报告里标注「达标 / 未达标」，
# 不参与退出码判定（见模块 docstring）。
TARGETS = {
    "冷启动（spawn → /health 200）": (5.0, "s"),
    "100 页文本 PDF 全流程": (480.0, "s"),
    "10 万字纯文本快速路径": (60.0, "s"),
    "引擎进程内存峰值": (4096.0, "MB"),
}


@dataclass
class Measurement:
    name: str
    value: float
    unit: str
    detail: str = ""
    target: float | None = None

    @property
    def verdict(self) -> str:
        if self.target is None:
            return "—"
        return "达标" if self.value <= self.target else "**未达标**"


@dataclass
class Report:
    rows: list[Measurement] = field(default_factory=list)

    def add(self, name: str, value: float, unit: str, detail: str = "") -> None:
        target = TARGETS.get(name, (None, unit))[0]
        self.rows.append(Measurement(name, value, unit, detail, target))
        mark = "" if target is None else ("  ✓" if value <= target else "  ✗ 超标")
        print(f"    {name}：{value:.2f} {unit}{mark}  {detail}", flush=True)


# ---------------------------------------------------------------- 样本生成


def make_100page_pdf(path: Path) -> int:
    """100 页数字 PDF，每页约 400 字正文。返回总字符数。"""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    sys.path.insert(0, str(ROOT / "corpus" / "fixtures"))
    from make_fixtures import _register_cjk_font

    font_name = _register_cjk_font()
    cn = font_name == "CJK"
    unit = ("本系统在完全离线的前提下完成文档结构化抽取，逐页记录解析级别与置信度。"
            if cn else
            "The system performs structured extraction fully offline, recording level and confidence per page. ")

    c = canvas.Canvas(str(path), pagesize=A4)
    _, height = A4
    total = 0
    for page in range(1, 101):
        c.setFont(font_name, 16)
        title = f"第 {page} 章 业务说明" if cn else f"Chapter {page} Business Notes"
        c.drawString(60, height - 70, title)
        total += len(title)
        c.setFont(font_name, 10.5)
        y = height - 105
        for _ in range(22):
            c.drawString(60, y, unit)
            total += len(unit)
            y -= 16
        c.showPage()
    c.save()
    return total


def make_100k_char_docx(path: Path) -> int:
    """约 10 万字的纯文本 docx（快速路径的载体）。返回总字符数。"""
    from docx import Document

    doc = Document()
    unit = "本条款约定双方的权利义务，任何一方不得单方面变更或解除，违约方应承担赔偿责任。"
    total = 0
    doc.add_heading("综合业务合同汇编", level=1)
    total += 8
    # 每段 ~40 字 × 2500 段 ≈ 10 万字
    for i in range(2500):
        if i % 100 == 0:
            h = f"第 {i // 100 + 1} 部分"
            doc.add_heading(h, level=2)
            total += len(h)
        doc.add_paragraph(unit)
        total += len(unit)
    doc.save(path)
    return total


# ---------------------------------------------------------------- 引擎驱动


class Engine:
    """拉起引擎子进程并提供 HTTP 客户端；退出时保证进程被清掉。"""

    def __init__(self, exe: str | None, data_dir: Path) -> None:
        self.exe = exe
        self.data_dir = data_dir
        self.token = secrets.token_hex(16)
        self.proc: subprocess.Popen[str] | None = None
        self.port = 0
        self.client: httpx.Client | None = None
        self.cold_start_s = 0.0
        self.peak_mem_mb = 0.0

    def __enter__(self) -> Engine:
        common = ["--port", "0", "--token", self.token, "--data-dir", str(self.data_dir)]
        if self.exe:
            path = Path(self.exe).resolve()
            cmd, cwd = [str(path), *common], str(path.parent)
        else:
            cmd, cwd = [sys.executable, "-m", "docfactory.main", *common], str(ENGINE_DIR)

        t0 = time.perf_counter()
        self.proc = subprocess.Popen(
            cmd, cwd=cwd, env={**os.environ, "PYTHONIOENCODING": "utf-8"},
            text=True, encoding="utf-8", errors="replace",
            stdout=subprocess.PIPE, stderr=subprocess.PIPE,
        )
        self.port = self._wait_ready()
        self.client = httpx.Client(
            base_url=f"http://127.0.0.1:{self.port}", timeout=120.0, trust_env=False,
            headers={"Authorization": f"Bearer {self.token}"},
        )
        # 冷启动口径取到 /health 200 为止：端口绑定不等于能服务（与主进程侧的判定一致）
        while True:
            try:
                if self.client.get("/health").status_code == 200:
                    break
            except httpx.HTTPError:
                pass
            if time.perf_counter() - t0 > READY_TIMEOUT_S:
                raise RuntimeError("引擎 /health 探活超时")
            time.sleep(0.05)
        self.cold_start_s = time.perf_counter() - t0
        return self

    def __exit__(self, *exc: object) -> None:
        if self.client is not None:
            with suppress(httpx.HTTPError):
                self.client.post("/shutdown")
            self.client.close()
        if self.proc is not None:
            try:
                self.proc.wait(timeout=20)
            except subprocess.TimeoutExpired:
                self.proc.kill()
                self.proc.wait(timeout=10)

    def _wait_ready(self) -> int:
        assert self.proc is not None
        deadline = time.monotonic() + READY_TIMEOUT_S
        while time.monotonic() < deadline:
            if self.proc.poll() is not None:
                err = self.proc.stderr.read() if self.proc.stderr else ""
                raise RuntimeError(f"引擎在握手前退出（{self.proc.returncode}）：{err}")
            line = (self.proc.stdout.readline() if self.proc.stdout else "").strip()
            if line.startswith("READY "):
                return int(json.loads(line[len("READY "):])["port"])
            if not line:
                time.sleep(0.05)
        raise RuntimeError("未在超时前收到 READY 握手")

    def sample_memory(self) -> None:
        """采一次引擎进程的工作集。用 tasklist 而不是 psutil：后者不是本工程依赖，
        为了测个内存把它拉进来不划算，而 tasklist 在目标平台上一定存在。"""
        if self.proc is None or sys.platform != "win32":
            return
        try:
            out = subprocess.run(
                ["tasklist", "/FI", f"PID eq {self.proc.pid}", "/FO", "CSV", "/NH"],
                capture_output=True, text=True, timeout=10,
                creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
            ).stdout
        except (OSError, subprocess.SubprocessError):
            return
        parts = [p.strip('" ') for p in out.strip().split('","')]
        if len(parts) < 5:
            return
        digits = "".join(ch for ch in parts[4] if ch.isdigit())
        if digits:
            self.peak_mem_mb = max(self.peak_mem_mb, int(digits) / 1024)

    def run_doc(self, sample: Path) -> tuple[float, dict]:
        """导入 + 解析一份文档，返回 (耗时秒, 文档详情)。期间持续采样内存。"""
        assert self.client is not None
        t0 = time.perf_counter()
        imported = self.client.post(
            "/documents/import", json={"paths": [str(sample)]}
        ).raise_for_status().json()
        entry = imported["imported"][0]
        doc_id, task_id = entry["doc_id"], entry["task_id"]

        last_sample = 0.0
        while True:
            detail = self.client.get(f"/tasks/{task_id}").raise_for_status().json()
            if detail["status"] in ("done", "failed", "canceled", "interrupted"):
                break
            now = time.perf_counter()
            if now - last_sample >= MEM_SAMPLE_INTERVAL_S:
                self.sample_memory()
                last_sample = now
            if now - t0 > TASK_TIMEOUT_S:
                raise RuntimeError(f"任务 {task_id} 超过 {TASK_TIMEOUT_S}s 未结束")
            time.sleep(0.1)
        elapsed = time.perf_counter() - t0
        self.sample_memory()
        if detail["status"] != "done":
            raise RuntimeError(f"解析未成功：{detail['status']} / {detail.get('error_code')}")
        doc = self.client.get(f"/documents/{doc_id}").raise_for_status().json()
        return elapsed, doc


# ---------------------------------------------------------------- 报告


def write_report(report: Report, path: Path, mode: str, sizes: dict[str, float]) -> None:
    lines = [
        "# 性能与体积基线（M1 实测）",
        "",
        f"- 运行模式：{mode}",
        f"- 平台：{sys.platform}　Python {sys.version.split()[0]}",
        "- 口径：真实进程 + HTTP 链路，非进程内直调",
        "",
        "> 08 章 §3.2 的门禁值要在 M2 依据本表冻结。**当前表中的「达标」只说明"
        "在本机这一次跑赢了目标值**，不构成验收结论——基准机（8 核 /16GB / 无 GPU）"
        "上的复测才算数。",
        "",
        "## 性能",
        "",
        "| 指标 | 实测 | 目标（01 章 §2.2） | 判定 | 备注 |",
        "|---|---|---|---|---|",
    ]
    for r in report.rows:
        target = "—" if r.target is None else f"≤ {r.target:g} {r.unit}"
        lines.append(f"| {r.name} | {r.value:.2f} {r.unit} | {target} | {r.verdict} | {r.detail} |")

    lines += [
        "",
        "## 体积（实测，未接入 M2 依赖）",
        "",
        "| 项 | 实测 | 08 章 §2 预估 |",
        "|---|---|---|",
    ]
    for name, mb in sizes.items():
        lines.append(f"| {name} | {mb:.0f} MB | {'—'} |")
    lines += [
        "",
        "**尚未计入**：Docling 模型（~250MB）、onnxruntime（~40MB）、OCR 模型（~20MB）、"
        "裁剪版 LibreOffice（~300MB）、Qwen tokenizer（~30MB）。",
        "这几项在 M2 接入后才会进包，届时需重跑本脚本并修订 08 章 §2 的体积清单。",
        "",
    ]
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(f"\n报告已写出：{path}")


def measure_sizes() -> dict[str, float]:
    """量已有产物的体积（缺失的项跳过，不报错）。"""
    def dir_mb(p: Path) -> float:
        if not p.exists():
            return 0.0
        return sum(f.stat().st_size for f in p.rglob("*") if f.is_file()) / 1048576

    out: dict[str, float] = {}
    engine_dist = ENGINE_DIR / "dist" / "engine"
    if engine_dist.exists():
        out["引擎 onedir（PyInstaller）"] = dir_mb(engine_dist)
    unpacked = ROOT / "app" / "dist" / "win-unpacked"
    if unpacked.exists():
        out["安装后总体积（win-unpacked）"] = dir_mb(unpacked)
    for setup in (ROOT / "app" / "dist").glob("*-setup.exe"):
        out[f"NSIS 安装包（{setup.name}）"] = setup.stat().st_size / 1048576
        break
    return out


def main(argv: list[str] | None = None) -> int:
    for stream in (sys.stdout, sys.stderr):
        if hasattr(stream, "reconfigure"):
            stream.reconfigure(encoding="utf-8", errors="replace")

    parser = argparse.ArgumentParser(description="性能基线实测（M1 交付物）")
    parser.add_argument("--exe", default=None, help="改测 PyInstaller 打包产物")
    parser.add_argument("--report", default=None, help="写出 Markdown 报告到该路径")
    args = parser.parse_args(argv)

    workdir = Path(tempfile.mkdtemp(prefix="df-perf-"))
    samples = workdir / "samples"
    samples.mkdir(parents=True, exist_ok=True)
    report = Report()

    print("[1/4] 生成大样本（100 页 PDF / 10 万字 docx）", flush=True)
    pdf_path = samples / "perf_100page.pdf"
    docx_path = samples / "perf_100k.docx"
    t0 = time.perf_counter()
    pdf_chars = make_100page_pdf(pdf_path)
    docx_chars = make_100k_char_docx(docx_path)
    print(f"    PDF {pdf_chars:,} 字符 / {pdf_path.stat().st_size / 1048576:.1f} MB；"
          f"docx {docx_chars:,} 字符（生成耗时 {time.perf_counter() - t0:.1f}s）", flush=True)

    mode = f"打包产物（{args.exe}）" if args.exe else "源码"
    print(f"[2/4] 拉起引擎（{mode}）", flush=True)
    with Engine(args.exe, workdir / "data-root") as eng:
        report.add("冷启动（spawn → /health 200）", eng.cold_start_s, "s",
                   f"端口 {eng.port}")

        print("[3/4] 100 页文本 PDF 全流程", flush=True)
        elapsed, doc = eng.run_doc(pdf_path)
        report.add("100 页文本 PDF 全流程", elapsed, "s",
                   f"{doc['page_cnt']} 页 / 级别 {doc['parse_level']} / "
                   f"覆盖率 {doc['text_coverage']} / {elapsed / max(doc['page_cnt'], 1):.2f}s 每页")

        print("[4/4] 10 万字纯文本快速路径", flush=True)
        elapsed, doc = eng.run_doc(docx_path)
        report.add("10 万字纯文本快速路径", elapsed, "s",
                   f"{docx_chars:,} 字符 / 级别 {doc['parse_level']}")

        report.add("引擎进程内存峰值", eng.peak_mem_mb, "MB",
                   f"采样间隔 {MEM_SAMPLE_INTERVAL_S}s")

    sizes = measure_sizes()
    if sizes:
        print("\n体积实测：", flush=True)
        for name, mb in sizes.items():
            print(f"    {name}：{mb:.0f} MB", flush=True)

    if args.report:
        write_report(report, Path(args.report), mode, sizes)

    over = [r for r in report.rows if r.target is not None and r.value > r.target]
    if over:
        print(f"\n注意：{len(over)} 项超出 01 章 §2.2 的目标值 —— "
              "但本脚本不据此判失败（门禁值待 M2 冻结）：", flush=True)
        for r in over:
            print(f"  · {r.name}：{r.value:.2f} {r.unit} > {r.target:g} {r.unit}", flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
